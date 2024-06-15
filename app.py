from flask import Flask, render_template, url_for, request, redirect, session, flash
from flask_sqlalchemy import SQLAlchemy
from flask_admin import Admin, AdminIndexView, expose, BaseView
from flask_admin.contrib.sqla import ModelView
from flask_admin.menu import MenuLink
from flask_admin.contrib import fileadmin
from flask_login import UserMixin, LoginManager, login_user, login_required, logout_user, current_user
from sqlalchemy.orm.exc import NoResultFound
from flask import Blueprint
import re
from flask_wtf import FlaskForm
from wtforms import StringField, SubmitField, PasswordField
from wtforms.validators import DataRequired
from flask_bcrypt import Bcrypt
from flask_httpauth import HTTPBasicAuth
import os
from docx import Document
from io import BytesIO
from docx.shared import Pt
from flask import send_file
import random
from werkzeug.security import generate_password_hash, check_password_hash
import speech_recognition as sr
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
import time
from datetime import datetime
from sqlalchemy.orm import Session
from sqlalchemy.orm import sessionmaker
from sqlalchemy.orm import sessionmaker, scoped_session
from sqlalchemy import create_engine



app = Flask(__name__)
auth = HTTPBasicAuth()
database_path = os.path.join(app.instance_path, 'sait.db')
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{database_path}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SECRET_KEY'] = os.urandom(24)
db = SQLAlchemy(app)
bcrypt = Bcrypt(app)
admin = Admin(app, name='MyAdmin', template_mode='bootstrap3')
login_manager = LoginManager(app)
login_manager.login_view = 'login'  # Указываем страницу входа

# Директория для загрузки фотографий
UPLOAD_FOLDER = 'C:\D\учеба\Flask\stud_sait\static\images'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
# Получите сессию SQLAlchemy

# Создание движка и сессии
engine = create_engine(app.config['SQLALCHEMY_DATABASE_URI'])
session_factory = sessionmaker(bind=engine)
db_session = scoped_session(session_factory)

class Students(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=True)
    game_name = db.Column(db.String(100), nullable=True)
    img = db.Column(db.String(100), nullable=True)
    url = db.Column(db.String(100), nullable=True)
    course = db.Column(db.Integer, nullable=True)
    role = db.Column(db.Integer, nullable=True)
    password = db.Column(db.String(100), nullable=True)


    def __repr__(self):
        return '<Students %r>' % self.id

    def set_password(self, password):
        self.password = bcrypt.generate_password_hash(password).decode('utf-8')

    def is_authenticated(self):
        if current_user.is_authenticated:
            print("Пользователь аутентифицирован")
        else:
            print("Пользователь не аутентифицирован")
        return True  # Возвращайте True для аутентифицированных пользователей

    @property
    def user_role(self):
        return self.role




class studweb(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=True)
    img = db.Column(db.String(100), nullable=True)
    studies = db.Column(db.Text, nullable=True)
    programm = db.Column(db.Text, nullable=True)
    course = db.Column(db.String(100), nullable=True)
    faculty = db.Column(db.String(100), nullable=True)
    description= db.Column(db.String(100), nullable=True)
    contacts= db.Column(db.String(100), nullable=True)


class MyAdminView(ModelView):
    column_list = ('id', 'name', 'role','img' ,'studies' ,'programm' ,'course' ,'faculty', 'photo' )  # Список полей, которые вы хотите отображать

    def get_query(self):
            if current_user.is_authenticated:
                if current_user.role == 3:  # Если роль пользователя 3
                    student_id = current_user.id
                    return self.session.query(self.model).filter_by(id=student_id)
            return self.session.query(self.model)  # Возвращаем все записи, если пользователь не аутентифицирован или его роль не 3


class List_diploms(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    year = db.Column(db.Integer, nullable=True)
    name_student = db.Column(db.String(100), nullable=True)
    name_diplom = db.Column(db.String(100), nullable=True)
    name_teacher = db.Column(db.String(100), nullable=True)


class List_coursework(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    year = db.Column(db.String(100), nullable=True)
    subject =db.Column(db.String(100), nullable=True)
    name_student = db.Column(db.String(100), nullable=True)
    name_coursework = db.Column(db.String(100), nullable=True)
    name_teacher = db.Column(db.String(100), nullable=True)


class Courses(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=True)
    faculty = db.Column(db.String(100), nullable=True)
    course = db.Column(db.Integer, nullable=True)
    groups_img = db.Column(db.String(100), nullable=True)
    curator = db.Column(db.String(100), nullable=True)
    curators_description = db.Column(db.String(100), nullable=True)
    curator_img = db.Column(db.String(100), nullable=True)
    role = db.Column(db.Integer, nullable=True)
    password = db.Column(db.String(100), nullable=True)

    def set_password(self, password):
        self.password = bcrypt.generate_password_hash(password).decode('utf-8')

    def is_authenticated(self):
        if current_user.is_authenticated:
            print("Пользователь аутентифицирован")
        else:
            print("Пользователь не аутентифицирован")
        return True  # Возвращайте True для аутентифицированных пользователей

    @property
    def user_role(self):
        return self.role

class Games(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    path = db.Column(db.String(100), nullable=True)
    img = db.Column(db.String(100), nullable=True)
    course = db.Column(db.Integer, nullable=True)
    name = db.Column(db.String(100), nullable=True)

class programm_achievements(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=True)
    programm_achievements = db.Column(db.String(100), nullable=True)
    description = db.Column(db.String(100), nullable=True)
    img = db.Column(db.String(100), nullable=True)
    link = db.Column(db.String(100), nullable=True)


class academic_achievements(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=True)
    academic_achievements = db.Column(db.String(100), nullable=True)
    description = db.Column(db.String(100), nullable=True)
    img = db.Column(db.String(100), nullable=True)
    link = db.Column(db.String(100), nullable=True)
    date = db.Column(db.String(100), nullable=True)


class characteristic(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=True)
    gender = db.Column(db.String(100), nullable=True)
    year_of_birth = db.Column(db.String(100), nullable=True)
    native = db.Column(db.String(100), nullable=True)
    city = db.Column(db.String(100), nullable=True)
    achieve = db.Column(db.String(100), nullable=True)
    departament = db.Column(db.String(100), nullable=True)
    recomend = db.Column(db.String(100), nullable=True)
    studies = db.Column(db.String(100), nullable=True)
    gpa = db.Column(db.String(100), nullable=True)
    number = db.Column(db.Integer, nullable=True)
    name_work = db.Column(db.String(100), nullable=True)
    form = db.Column(db.String(100), nullable=True)
    output = db.Column(db.String(100), nullable=True)
    pages = db.Column(db.String(100), nullable=True)
    co_authors = db.Column(db.String(100), nullable=True)
    subject = db.Column(db.String(100), nullable=True)
    ecolog = db.Column(db.String(100), nullable=True)
    patriot = db.Column(db.String(100), nullable=True)
    family = db.Column(db.String(100), nullable=True)
    character = db.Column(db.String(100), nullable=True)
    authority = db.Column(db.String(100), nullable=True)
    shows = db.Column(db.String(100), nullable=True)
    violations = db.Column(db.String(100), nullable=True)


class curators_report(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    job_title = db.Column(db.String(100), nullable=True)
    half_year = db.Column(db.String(100), nullable=True)
    year = db.Column(db.String(100), nullable=True)
    month = db.Column(db.String(100), nullable=True)
    group = db.Column(db.String(100), nullable=True)
    curator = db.Column(db.String(100), nullable=True)
    starosta = db.Column(db.String(100), nullable=True)
    proforg = db.Column(db.String(100), nullable=True)
    studcom = db.Column(db.String(100), nullable=True)
    curators_hour = db.Column(db.String(), nullable=True)
    date = db.Column(db.String(100), nullable=True)
    percent = db.Column(db.String(100), nullable=True)
    grajd_patriot = db.Column(db.String(), nullable=True)
    nravst_estet = db.Column(db.String(), nullable=True)
    trad_vuz = db.Column(db.String(), nullable=True)
    zdorov_obraz = db.Column(db.String(), nullable=True)
    prof_mer = db.Column(db.String(), nullable=True)
    trud_vosp = db.Column(db.String(), nullable=True)


class Contact(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=True)
    message = db.Column(db.Text, nullable=True)

class BullCowUser(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    player_name = db.Column(db.String(50), nullable=True)
    number = db.Column(db.Integer, nullable=True)
    result = db.Column(db.String(50), nullable=True)

class guess_num_user(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    player_name = db.Column(db.String(50), nullable=True)
    number = db.Column(db.Integer, nullable=True)
    result = db.Column(db.String(50), nullable=True)
    attempts = db.Column(db.String(50), nullable=True)




class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True)
    password = db.Column(db.String(60))
    role = db.Column(db.Integer, nullable=True)


class Roles(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), unique=True)


class MyCustomView2(BaseView):
    @expose('/')
    def index(self):
        # Получаем данные из таблицы courses
        courses_data = Courses.query.all()
        return render_template('admin/my_custom_button.html', courses_data=courses_data, endpoint='mycustomview2')


    @expose('/students/<group_name>')
    def students_by_group(self, group_name):
        # Находим номер курса по имени группы
        course = Courses.query.filter_by(name=group_name).first()
        if not course:
            return "Group not found", 404

        course_number = course.course

        # Получаем студентов по номеру курса
        students_data = studweb.query.filter_by(course=course_number).all()

        # Возвращаем шаблон с данными
        return render_template('admin/students_by_group.html', students_data=students_data, group_name=group_name)

    @expose('/create_characteristic', methods=['POST'])
    def create_characteristic(self):
        id = request.form.get('student_id')
        if not id:
            flash('No student selected', 'error')
            return redirect(url_for('mycustomview.index'))

        # Использование сессии для выполнения запросов
        with db_session() as session:
            stud = session.query(Students).get(id)
            print(stud)
            # Получение имени студента
            student_name = stud.name
            print(student_name)
            student = session.query(characteristic).get(id)
            # Получение академических достижений по имени студента
            # Получение академических достижений студента из базы данных
            academic_achievementss = session.query(academic_achievements).filter_by(name=student_name).all()
            print(academic_achievementss)
            # Вывод всех академических достижений студента
            for achievement in academic_achievementss:
                print("Тема:", achievement.academic_achievements)
                print("Курс:", achievement.img)
                print("Описание:", achievement.description)
                print()
            programm_achievementss = session.query(programm_achievements).filter_by(name=student_name).all()
            # Вывод всех академических достижений студента
            for achievement in programm_achievementss:
                print("Тема:", achievement.programm_achievements)
                print("Курс:", achievement.img)
                print("Описание:", achievement.description)
                print()
            if not student:
                flash('Student not found', 'error')
                return redirect(url_for('mycustomview.index'))

        # Получение данных из базы данных
        student_name = student.name
        student_gender = student.gender
        student_year_of_birth = student.year_of_birth
        student_native = student.native
        student_city = student.city
        course_number = stud.course
        student_achieve= student.achieve

        course = Courses.query.filter_by(course=course_number).first()
        if not course:
            flash('Course not found', 'error')
            return redirect(url_for('mycustomview.index'))

        group_name = course.name

        # Путь к шаблонному документу
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')

        # Открытие шаблонного документа
        try:
            template_doc = Document(template_path)
        except Exception as e:
            flash(f'Error loading template: {e}', 'error')
            return redirect(url_for('mycustomview.index'))

        student_middle_name = student_name.split()[1] if len(student_name.split()) > 1 else student_name

        # Создание строки с академическими достижениями
        academic_achievements_str = "\n".join([f" {achievement.academic_achievements}" for achievement in academic_achievementss])

        # Создание строки с программными достижениями
        programm_achievements_str = "\n".join([f" {achievement.programm_achievements}" for achievement in programm_achievementss])

        # Подстановка данных в словарь
        replace_dict = {
            "{GENDER}": student_gender,
            "{COURSE_NUMBER}": arabic_to_roman(course_number),
            "{COURSE_GROUP}": group_name,
            "{STUDENT_FIO}": student_name,
            "{YEAR_OF_BIRTH}": str(student_year_of_birth),
            "{NATIVE}": student_native,
            "{CITY}": student_city,
            "{STUDENT_NAME}": student_middle_name,
            "{STUDENT_ACHIVE}": student_achieve,
            "{STUDENT_academic_achievements}": academic_achievements_str,
            "{STUDENT_programm_achievements}": programm_achievements_str
        }

        # Отладочная информация
        print("Замена данных:")
        for key, value in replace_dict.items():
            print(f"{key}: {value}")

       # Подстановка данных в шаблонный документ
        for paragraph in template_doc.paragraphs:
            for key, value in replace_dict.items():
                if key in paragraph.text:
                    new_text = paragraph.text.replace(key, value)
                    paragraph.clear()
                    run = paragraph.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)


        # Сохранение нового документа в объект BytesIO
        output = BytesIO()
        template_doc.save(output)
        output.seek(0)
        # Отправка файла клиенту
        return send_file(output, as_attachment=True, download_name=f'Характеристика_{student_name}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    @expose('/create_group_report', methods=['POST'])
    def create_group_report(self):
        group_name = request.form.get('group_name')
        if not group_name:
            flash('No group selected', 'error')
            return redirect(url_for('mycustomview.index'))

        # Использование сессии для выполнения запросов
        with db_session() as session:
            #student = session.query(Students).get(id)
            report = curators_report.query.filter_by(group=group_name).first()
            print(report)
            if not report:
                flash('Student not found', 'error')
                return redirect(url_for('mycustomview.index'))

        job_title = report.job_title
        half_year = report.half_year
        year = report.year
        month = report.month
        group = report.group
        curator = report.curator
        starosta = report.starosta
        proforg = report.proforg
        studcom = report.studcom
        curators_hour = report.curators_hour
        dates = report.date.split()
        percents = report.percent.split()
        grajd_patriot = report.grajd_patriot.replace("\n", "\n")
        nravst_estet = report.nravst_estet.replace("\n", "\n")
        trad_vuz = report.trad_vuz.replace("\n", "\n")
        zdorov_obraz = report.zdorov_obraz.replace("\n", "\n")
        prof_mer = report.prof_mer.replace("\n", "\n")
        trud_vosp = report.trud_vosp.replace("\n", "\n")

        # Путь к шаблонному документу
        template_path = os.path.join(os.path.dirname(__file__), 'template_curator.docx')

        # Открытие шаблонного документа
        try:
            template_doc = Document(template_path)
        except Exception as e:
            flash(f'Error loading template: {e}', 'error')
            return redirect(url_for('mycustomview1.index'))

        # Разделение curators_hour на блоки
        curators_hour_blocks = []
        current_block = []
        for line in curators_hour.split("\n"):
            if line.startswith("1. ") and current_block:
                curators_hour_blocks.append("\n".join(current_block))
                current_block = [line]
            else:
                current_block.append(line)
        if current_block:
            curators_hour_blocks.append("\n".join(current_block))

        # Подстановка данных в словарь
        replace_dict = {
            "{HALF_YEAR}": half_year,
            "{YEAR}": year,
            "{MONTH}": month,
            "{GROUP_NUMBER}": group,
            "{CURATOR}": curator,
            "{STAROSTA}": starosta,
            "{PROFORG}": proforg,
            "{STUDCOM}": studcom,
            "{GRAJD_PATRIOT}": grajd_patriot,
            "{TRAD_VUZ}": trad_vuz,
            "{NRAVST_ESTET}": nravst_estet,
            "{ZDOROV_OBRAZ}": zdorov_obraz,
            "{PROF_MER}": prof_mer,
            "{TRUD_VOSP}": trud_vosp,
            "{JOB_TITLE}": job_title
        }

        # Отладочная информация
        print("Замена данных:")
        for key, value in replace_dict.items():
            print(f"{key}: {value}")

        # Подстановка данных в шаблонный документ
        for paragraph in template_doc.paragraphs:
            for key, value in replace_dict.items():
                if key in paragraph.text:
                    new_text = paragraph.text.replace(key, value)
                    paragraph.clear()
                    run = paragraph.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

        # Обработка кураторских часов
        for paragraph in template_doc.paragraphs:
            if "{CURATORS_HOUR}" in paragraph.text:
                for i, block in enumerate(curators_hour_blocks):
                    new_text = f"{dates[i]} – {percents[i]} %\n\n{block}"
                    # Вставка новой секции перед текущим параграфом
                    new_paragraph = paragraph.insert_paragraph_before(new_text)
                    for run in new_paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                # Удаление заполнительного текста
                paragraph.clear()

        # Сохранение нового документа в объект BytesIO
        output = BytesIO()
        template_doc.save(output)
        output.seek(0)

        # Отправка файла клиенту
        return send_file(output, as_attachment=True, download_name=f'Отчет_группы_{group}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        return self.students_by_group(group_name)
def arabic_to_roman(n):
    roman_numerals = {
        1: 'I', 2: 'II', 3: 'III', 4: 'IV', 5: 'V',
        6: 'VI', 7: 'VII', 8: 'VIII', 9: 'IX', 10: 'X'
    }
    return roman_numerals.get(n, str(n))


class MyCustomView1(BaseView):
    @expose('/')
    def index(self):
        # Получаем данные из таблицы courses
        courses_data = Courses.query.all()
        return render_template('admin/my_custom_button.html', courses_data=courses_data, endpoint='mycustomview1')

    @expose('/students/<group_name>')
    def students_by_group(self, group_name):
        # Находим номер курса по имени группы
        course = Courses.query.filter_by(name=group_name).first()
        if not course:
            return "Group not found", 404

        course_number = course.course

        # Получаем студентов по номеру курса
        students_data = studweb.query.filter_by(course=course_number).all()

        # Возвращаем шаблон с данными
        return render_template('admin/students_by_group.html', students_data=students_data, group_name=group_name, endpoint='mycustomview1')

    @expose('/create_characteristic', methods=['POST'])
    def create_characteristic(self):
        student_id = request.form.get('student_id')
        if not student_id:
            flash('No student selected', 'error')
            return redirect(url_for('mycustomview1.index'))

        # Использование сессии для выполнения запросов
        with db_session() as session:
            stud = session.query(Students).get(student_id)
            student_name = stud.name
            student = session.query(characteristic).get(student_id)
            academic_achievementss = session.query(academic_achievements).filter_by(name=student_name).all()
            programm_achievementss = session.query(programm_achievements).filter_by(name=student_name).all()
            if not student:
                flash('Student not found', 'error')
                return redirect(url_for('mycustomview1.index'))

        # Получение данных из базы данных
        student_name = student.name
        student_gender = student.gender
        student_year_of_birth = student.year_of_birth
        student_native = student.native
        student_city = student.city
        course_number = stud.course
        student_achieve= student.achieve

        course = Courses.query.filter_by(course=course_number).first()
        if not course:
            flash('Course not found', 'error')
            return redirect(url_for('mycustomview.index'))

        group_name = course.name

        # Путь к шаблонному документу
        template_path = os.path.join(os.path.dirname(__file__), 'template.docx')

        # Открытие шаблонного документа
        try:
            template_doc = Document(template_path)
        except Exception as e:
            flash(f'Error loading template: {e}', 'error')
            return redirect(url_for('mycustomview.index'))

        student_middle_name = student_name.split()[1] if len(student_name.split()) > 1 else student_name

        # Создание строки с академическими достижениями
        academic_achievements_str = "\n".join([f" {achievement.academic_achievements}" for achievement in academic_achievementss])

        # Создание строки с программными достижениями
        programm_achievements_str = "\n".join([f" {achievement.programm_achievements}" for achievement in programm_achievementss])

        # Подстановка данных в словарь
        replace_dict = {
            "{GENDER}": student_gender,
            "{COURSE_NUMBER}": arabic_to_roman(course_number),
            "{COURSE_GROUP}": group_name,
            "{STUDENT_FIO}": student_name,
            "{YEAR_OF_BIRTH}": str(student_year_of_birth),
            "{NATIVE}": student_native,
            "{CITY}": student_city,
            "{STUDENT_NAME}": student_middle_name,
            "{STUDENT_ACHIVE}": student_achieve,
            "{STUDENT_academic_achievements}": academic_achievements_str,
            "{STUDENT_programm_achievements}": programm_achievements_str
        }

        # Отладочная информация
        print("Замена данных:")
        for key, value in replace_dict.items():
            print(f"{key}: {value}")

       # Подстановка данных в шаблонный документ
        for paragraph in template_doc.paragraphs:
            for key, value in replace_dict.items():
                if key in paragraph.text:
                    new_text = paragraph.text.replace(key, value)
                    paragraph.clear()
                    run = paragraph.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)


        # Сохранение нового документа в объект BytesIO
        output = BytesIO()
        template_doc.save(output)
        output.seek(0)

        # Отправка файла клиенту
        return send_file(output, as_attachment=True, download_name=f'Характеристика_{student_name}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


    @expose('/create_group_report', methods=['POST'])
    def create_group_report(self):
        group_name = request.form.get('group_name')
        if not group_name:
            flash('No group selected', 'error')
            return redirect(url_for('mycustomview.index'))

        # Использование сессии для выполнения запросов
        with db_session() as session:
            #student = session.query(Students).get(id)
            report = curators_report.query.filter_by(group=group_name).first()
            print(report)
            if not report:
                flash('Student not found', 'error')
                return redirect(url_for('mycustomview.index'))

        job_title = report.job_title
        half_year = report.half_year
        year = report.year
        month = report.month
        group = report.group
        curator = report.curator
        starosta = report.starosta
        proforg = report.proforg
        studcom = report.studcom
        curators_hour = report.curators_hour
        dates = report.date.split()
        percents = report.percent.split()
        grajd_patriot = report.grajd_patriot.replace("\n", "\n") if report.grajd_patriot else ""
        nravst_estet = report.nravst_estet.replace("\n", "\n") if report.nravst_estet else ""
        trad_vuz = report.trad_vuz.replace("\n", "\n") if report.trad_vuz else ""
        zdorov_obraz = report.zdorov_obraz.replace("\n", "\n") if report.zdorov_obraz else ""
        prof_mer = report.prof_mer.replace("\n", "\n") if report.prof_mer else ""
        trud_vosp = report.trud_vosp.replace("\n", "\n") if report.trud_vosp else ""


        # Путь к шаблонному документу
        template_path = os.path.join(os.path.dirname(__file__), 'template_curator.docx')

        # Открытие шаблонного документа
        try:
            template_doc = Document(template_path)
        except Exception as e:
            flash(f'Error loading template: {e}', 'error')
            return redirect(url_for('mycustomview1.index'))

        # Разделение curators_hour на блоки
        curators_hour_blocks = []
        current_block = []
        for line in curators_hour.split("\n"):
            if line.startswith("1. ") and current_block:
                curators_hour_blocks.append("\n".join(current_block))
                current_block = [line]
            else:
                current_block.append(line)
        if current_block:
            curators_hour_blocks.append("\n".join(current_block))

        # Подстановка данных в словарь
        replace_dict = {
            "{HALF_YEAR}": half_year,
            "{YEAR}": year,
            "{MONTH}": month,
            "{GROUP_NUMBER}": group,
            "{CURATOR}": curator,
            "{STAROSTA}": starosta,
            "{PROFORG}": proforg,
            "{STUDCOM}": studcom,
            "{GRAJD_PATRIOT}": grajd_patriot,
            "{TRAD_VUZ}": trad_vuz,
            "{NRAVST_ESTET}": nravst_estet,
            "{ZDOROV_OBRAZ}": zdorov_obraz,
            "{PROF_MER}": prof_mer,
            "{TRUD_VOSP}": trud_vosp,
            "{JOB_TITLE}": job_title
        }

        # Отладочная информация
        print("Замена данных:")
        for key, value in replace_dict.items():
            print(f"{key}: {value}")

        # Подстановка данных в шаблонный документ
        for paragraph in template_doc.paragraphs:
            for key, value in replace_dict.items():
                if key in paragraph.text:
                    new_text = paragraph.text.replace(key, value)
                    paragraph.clear()
                    run = paragraph.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

        # Обработка кураторских часов
        for paragraph in template_doc.paragraphs:
            if "{CURATORS_HOUR}" in paragraph.text:
                for i, block in enumerate(curators_hour_blocks):
                    new_text = f"{dates[i]} – {percents[i]} %\n\n{block}"
                    # Вставка новой секции перед текущим параграфом
                    new_paragraph = paragraph.insert_paragraph_before(new_text)
                    for run in new_paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)
                # Удаление заполнительного текста
                paragraph.clear()

        # Сохранение нового документа в объект BytesIO
        output = BytesIO()
        template_doc.save(output)
        output.seek(0)

        # Отправка файла клиенту
        return send_file(output, as_attachment=True, download_name=f'Отчет_группы_{group}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        return self.students_by_group(group_name)
    def arabic_to_roman(n):
        roman_numerals = {
            1: 'I', 2: 'II', 3: 'III', 4: 'IV', 5: 'V',
            6: 'VI', 7: 'VII', 8: 'VIII', 9: 'IX', 10: 'X'
        }
        return roman_numerals.get(n, str(n))



class MyCustomView3(BaseView):
    @expose('/')
    @expose('/form_kursovaya', methods=['GET'])
    def form_kursovaya(self):
        return render_template('admin/form_kursovaya.html')

    @expose('/create_kursovaya', methods=['POST'])
    def create_work(self):
        work_type = request.form.get('work_type')
        subject = request.form.get('subject')
        theme = request.form.get('theme')
        gender = request.form.get('gender')
        course_number = request.form.get('course_number')
        student_fio = request.form.get('student_fio')
        title = request.form.get('title')
        teacher_fio = request.form.get('teacher_fio')
        year = request.form.get('year')


        if work_type == 'kursovaya':
            template_path = 'template_kursovaya.docx'
            file_name = f'{student_fio}_kursovaya.docx'
        else:
            template_path = 'template_diplom.docx'
            file_name = f'{student_fio}_diploma.docx'

        # Подстановка данных в словарь
        replace_dict = {
            "{SUBJECT}": subject,
            "{THEME}": theme,
            "{GENDER}": gender,
            "{COURSE_NUMBER}": course_number,
            "{STUDENT_FIO}": student_fio,
            "{TITLE}": title,
            "{TEACHER_FIO}": teacher_fio,
            "{YEAR}": year
        }

        doc = Document(template_path)

        # Подстановка данных в шаблонный документ
        for paragraph in doc.paragraphs:
            for key, value in replace_dict.items():
                if key in paragraph.text:
                    new_text = paragraph.text.replace(key, value)
                    paragraph.clear()
                    run = paragraph.add_run(new_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)

        # Save the filled document to a BytesIO object
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        # Send the file to the user
        return send_file(output, as_attachment=True, download_name=file_name, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')




admin.add_view(ModelView(User, db.session))
admin.add_view(ModelView(Students, db.session))
admin.add_view(ModelView(Courses, db.session))
admin.add_view(ModelView(Games, db.session))
admin.add_view(ModelView(Contact, db.session))
admin.add_view(ModelView(academic_achievements, db.session))
admin.add_view(ModelView(programm_achievements, db.session))
admin.add_view(MyAdminView(studweb, db.session, name='studweb'))
admin.add_view(ModelView(curators_report, db.session))
admin.add_view(ModelView(characteristic, db.session))
admin.add_view(MyCustomView1(name='teacher', endpoint='mycustomview1'))
admin.add_view(MyCustomView2(name='zav', endpoint='mycustomview2'))
admin.add_view(MyCustomView3(name='Create Work', endpoint='mycustomview3'))

admin_bp = Blueprint('MyAdmin', __name__, url_prefix='/admin')
admin.init_app(admin_bp)


@app.route('/hash_password/<int:student_id>', methods=['GET', 'POST'])
def hash_password(student_id):
    student = Students.query.get(student_id)

    if student:
        if request.method == 'POST':
            new_password = request.form['new_password']
            student.set_password(new_password)
            db.session.commit()
            flash(f"Пароль студента {student_id} успешно хеширован.", 'success')
            return redirect(url_for('index'))  # Замените 'some_route' на ваш роут

        return render_template('change_password.html', student=student)
    else:
        return f"Студент с ID {student_id} не найден."


@app.route('/hash_password_teacher/<int:teacher_id>', methods=['GET', 'POST'])
def hash_password_teacher(teacher_id):
    teacher = Courses.query.get(teacher_id)

    if teacher:
        if request.method == 'POST':
            new_password = request.form['new_password']
            teacher.set_password(new_password)
            db.session.commit()
            flash(f"Пароль преподавателя {teacher_id} успешно хеширован.", 'success')
            return redirect(url_for('index'))  # Замените 'index' на ваш роут

        return render_template('change_password_teacher.html', teacher=teacher)
    else:
        return f"Преподаватель с ID {teacher_id} не найден."

@auth.verify_password
def verify_password(username, password):
    student = Students.query.filter_by(name=username).first()
    if student and bcrypt.check_password_hash(student.password, password):
        return (True, student.role)
    return False


admin.add_link(MenuLink(name='Logout', category='', url="/logout"))
#admin.add_link(MenuLink(name='Login', category='', url="/login"))

class MyAdminView(BaseView):
    @expose('/')
    def index_view(self):
        student_id = request.args.get('student_id')


# Функция загрузки пользователя для Flask-Login
@login_manager.user_loader
def load_user(user_id):
    with db_session() as session:
        student = session.query(Students).get(user_id)
    return student

@app.route('/login', methods=['GET', 'POST'])
def login():
    # Добавьте проверку, если пользователь уже вошел
    if current_user.is_authenticated:
        flash("Вы уже вошли в систему.", 'info')
        return redirect(url_for('admin.index')) # Предполагаем, что это маршрут админ-панели

    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        student = Students.query.filter_by(name=username).first()
        teacher = Courses.query.filter_by(name=username).first()

        if student and bcrypt.check_password_hash(student.password, password):
            user_instance = Students.query.filter_by(name=username).first()
            if not user_instance:
                user_instance = Students(name=username, password=generate_password_hash(password), role=student.role)
                db.session.add(user_instance)
                db.session.commit()
            login_user(user_instance)
            flash("Вход выполнен успешно", 'success')
            if student.role == 1:
                return redirect(url_for('admin.index'))
            elif student.role == 2:
                return redirect(url_for('students.index_view'))
            elif student.role == 3:
                return redirect(url_for('studweb.index_view', student_id=student.id))

        elif teacher and bcrypt.check_password_hash(teacher.password, password):
            user_instance = Courses.query.filter_by(name=username).first()
            if not user_instance:
                user_instance = Courses(name=username, password=generate_password_hash(password), role=teacher.role)
                db.session.add(user_instance)
                db.session.commit()
            login_user(user_instance)
            flash("Вход выполнен успешно", 'success')
            # Вероятно, у вас есть разные маршруты для преподавателей и студентов,
            # замените соответствующие маршруты ниже
            if teacher.role == 5:
                 return redirect(url_for('mycustomview2.index'))
            elif teacher.role == 4:
                return redirect(url_for('mycustomview1.students_by_group', group_name=username ))

        else:
            flash("Неправильное имя пользователя или пароль", 'error')

    return render_template('login.html')


@app.route('/logout/')
@login_required
def logout():
    logout_user()
    flash("Вы вышли из системы.")
    return redirect(url_for('login'))


# Маршрут админки
@app.route('/admin')
@login_required
def admin():
    if current_user.role != 1:  # Если у пользователя нет админских прав
        flash("У вас нет доступа к этой странице.", 'error')
        return redirect(url_for('login'))

    return redirect(url_for('admin'))


def generate_secret_number():
    # Генерируем случайное число из четырех цифр
    first_digit = random.randint(1, 9)  # Генерируем первую цифру от 1 до 9
    rest_of_digits = ''.join([str(random.randint(0, 9)) for _ in range(3)])  # Генерируем остальные три цифры
    secret_number = str(first_digit) + rest_of_digits
    return secret_number


@app.route('/')
@app.route('/home')
def index():
    students = Students.query.order_by(Students.id.asc()).all()
    # Извлекаем все группы из базы данных
    groups = Courses.query.order_by(Courses.id.asc()).all()
    # Передаем группы в шаблон
    return render_template("index.html", groups=groups)


@app.route('/ListOfGroups')
def ListOfGroups():
    # Извлекаем все группы из базы данных
    groups = Courses.query.order_by(Courses.id.asc()).all()

    return render_template("ListOfGroups.html", groups=groups)

@app.route('/Contacts', methods=['GET', 'POST'])
def Contacts():
    if request.method == 'POST':
        name = request.form.get('name')
        message = request.form.get('message')

        if name or message:
            new_contact = Contact(name=name, message=message)
            db.session.add(new_contact)
            db.session.commit()
            flash('Your message has been sent!', 'success')
        else:
            flash('Please fill in at least one field.', 'error')

        return redirect(url_for('Contacts'))

    return render_template('Contacts.html')
@app.route('/base')
def base():
    return render_template("base.html",)

@app.route('/course<int:id>', methods=['POST', 'GET'] )
def course(id):
    courses = Courses.query.order_by(Courses.id.asc()).all()  # по id вывести группу
    #print(courses)
    course = Courses.query.filter_by(id=id).first()

    #print(course.id)
    #for el in courses:
        #print(el.course)
    groups = Courses.query.with_entities(Courses.id, Courses.name).all()
    #print(groups)
    students = Students.query.order_by(Students.id.asc()).all()
    #for el in students:
        #print(el.course)
    games = Courses.query.with_entities(Games.img, Games.name, Games.path).all()
    #print(games)
    # Извлекаем значения поля Name из объектов Courses
    names = [course.name for course in courses]

    return render_template("course_page.html", students=students, course=course, groups=groups, games=games, courses=courses )


@app.route('/student<int:id>',  methods=['POST', 'GET'] )
def students_web(id):
    try:
        student = db.session.query(studweb).filter_by(id=id).one()

        programm_sentences = programm_achievements.query.filter_by(name=student.name).all()
        academic_sentences = academic_achievements.query.filter_by(name=student.name).all()
        #studies_sentences = re.findall(r'([А-ЯA-Z].*?[.!?])', student.studies)
        groups = Courses.query.with_entities(Courses.id, Courses.name).all()

        return render_template("student_page.html", student=student, programm_sentences=programm_sentences, academic_sentences=academic_sentences, groups=groups)
    except NoResultFound:
        # Handle the case where the student with the given id is not found.
        # You can return an error page or a 404 Not Found response.
        return render_template("student_not_found.html", student=student)

@app.route('/Lists')
def List():
    return render_template("List.html")


#@app.route('/diplom')
#def Diplom():
#    return render_template("Diplom.html")


@app.route('/diplom')
def diplom_year():
    years = List_diploms.query.distinct(List_diploms.year).order_by(List_diploms.year).all()
    return render_template('Diplom_years.html', years=years)
@app.route('/diplom/<int:year>')
def diplom(year):
    data = List_diploms.query.filter_by(year=year).all()
    return render_template('Diplom.html', data=data, year=year)

from sqlalchemy import or_

@app.route('/search_teacher_diplom')
def search_teacher_diplom():
    teacher_name = request.args.get('teacher')
    if teacher_name:
        teachers = List_diploms.query.filter(or_(List_diploms.name_teacher.like(f"%{teacher_name}%"))).all()
        return render_template('teacher_results.html', teacher_name=teacher_name, teachers=teachers)
    else:
        # Обработка случая, когда пользователь не ввел фамилию преподавателя
        return render_template('error.html', message="Введите фамилию преподавателя для поиска")


@app.route('/kursovaya')
def kursovaya_year():
    years = List_coursework.query.distinct(List_coursework.year).order_by(List_coursework.year).all()
    return render_template('Kursovaya_years.html', years=years)
@app.route('/kursovaya/<int:year>')
def kursovaya(year):
    data = List_coursework.query.filter_by(year=year).all()
    return render_template('Kursovaya.html', data=data, year=year)

@app.route('/search_teacher_kursovaya')
def search_teacher_kursovaya():
    teacher_name = request.args.get('teacher')
    if teacher_name:
        teachers = List_coursework.query.filter(or_(List_coursework.name_teacher.like(f"%{teacher_name}%"))).all()
        return render_template('teacher_results_kursovaya.html', teacher_name=teacher_name, teachers=teachers)
    else:
        # Обработка случая, когда пользователь не ввел фамилию преподавателя
        return render_template('error.html', message="Введите фамилию преподавателя для поиска")

@app.route('/bull_cow_log', methods=['GET', 'POST'])
def bull_cow_log():
    if request.method == 'POST':
        player_name = request.form.get('playerName')
        number = request.form.get('guessNumber')
        secret_number = generate_secret_number()
        session['secret_number'] = secret_number

        return redirect(url_for('bull_cow', player_name=player_name, number=number))
    return render_template('index_bull_cow.htm')

def generate_secret_number_guess():
    return random.randint(1, 100)

@app.route('/guess_num', methods=['GET', 'POST'])
def guess_num():
    if request.method == 'POST':
        player_name = request.form.get('playerName')
        guess_number = int(request.form.get('guessNumber'))

        if 'attempts' not in session:
            session['attempts'] = 0
            session['guesses'] = []

        session['attempts'] += 1
        session['guesses'].append(guess_number)
        if 'secret_number' not in session:
            session['secret_number'] = generate_secret_number_guess()

        secret_number = session['secret_number']
        print(secret_number)
        if session['attempts'] > 10:
            flash('Вы исчерпали все 10 попыток. Чтобы начать игру заново, введите имя и число и нажмите "Проверить".', 'error')
            result = 'Вы проиграли'
            game_result = guess_num_user(player_name=player_name, number=secret_number, result=result, attempts=session['attempts'])
            db.session.add(game_result)
            db.session.commit()
            session.pop('secret_number', None)
            session.pop('attempts', None)
            session.pop('guesses', None)
            return redirect(url_for('guess_num'))



        if int(guess_number) < int(secret_number):
            message = f"Ваше число {guess_number} меньше задуманного."
            result = 'Не угадал'
        elif int(guess_number) > int(secret_number):
            message = f"Ваше число {guess_number} больше задуманного."
            result = 'Не угадал'
        else:
            message = f"Поздравляем, {player_name}! Вы угадали число {secret_number}."
            result = 'Вы победили'
            game_result = guess_num_user(player_name=player_name, number=secret_number, result=result, attempts=session['attempts'])
            db.session.add(game_result)
            db.session.commit()
            session.pop('secret_number', None)
            session.pop('attempts', None)
            session.pop('guesses', None)
            flash(message)
            return redirect(url_for('guess_num', player_name=player_name))

        flash(message)
        return redirect(url_for('guess_num', player_name=player_name))

    player_name = request.args.get('player_name', '')
    attempts = session.get('attempts', 0)
    guesses = session.get('guesses', [])
    return render_template('index_end.htm', player_name=player_name, attempts=attempts, guesses=guesses)



@app.route('/bull_cow/<player_name>/<number>', methods=['GET'])
def bull_cow(player_name, number):
    #global attempts_count  # Объявляем, что мы будем использовать глобальный счетчик
    if 'attempts_count' not in session:
        session['attempts_count'] = 1
    else:
        session['attempts_count'] += 1

    attempts_count = session['attempts_count']
    if attempts_count > 2:
        session.pop('attempts_count')
        result = 'Вы проиграли!'
        return render_template('bull_cow_restart.html', player_name=player_name)

        # Дополнительные действия при достижении лимита попыток, если нужны
    secret_number = session.get('secret_number')
    print(secret_number)
    bulls, cows = check_guess(number, secret_number)
    result = 'Вы победили!'
    if bulls == 4:  # Все цифры на своих местах
        new_record = BullCowUser(player_name=player_name, number=number, result=result)
        db.session.add(new_record)
        db.session.commit()
        return render_template('bull_cow_congratulations.html', player_name=player_name)
    else:
        return render_template('bull_cow_try_again.html', player_name=player_name, number=number, bulls=bulls,
                               cows=cows, attempts_count=attempts_count)


def check_guess(guess, secret):
    bulls, cows = 0, 0
    for i in range(4):
        if guess[i] == secret[i]:
            bulls += 1
        elif guess[i] in secret:
            cows += 1
    return bulls, cows


@app.route('/about')
def about():
    groups = Courses.query.with_entities(Courses.id, Courses.name).all()
    return render_template("about.html", groups=groups)


if __name__ == "__main__":
    app.run(debug=True,port=5000)

# Проверка аутентификации и роли пользователя перед каждым запросом
@app.before_request
def restrict_access():
    if request.endpoint and "/admin" in request.endpoint:  # Проверяем, что запрос направлен к админке
        if not current_user.is_authenticated:  # Если пользователь не аутентифицирован
            return redirect(url_for('login'))  # Перенаправляем его на страницу входа
        elif current_user.user_role != 1:  # Если у пользователя нет админских прав
            flash("У вас нет доступа к этой странице.", 'error')  # Показываем сообщение об ошибке
            return redirect(url_for('login'))  # Перенаправляем его на страницу входа
